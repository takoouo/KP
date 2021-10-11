from typing import List
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *

import certificate as ui
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.table import _Cell
from docx2pdf import convert
import os

def input_word(t, x, y, w, a):
    p = t.cell(x,y).paragraphs[0]
    t.cell(x,y).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p.paragraph_format.alignment=a
    p.add_run(w).font.size = Pt(13)

def Set_cell_border(cell: _Cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def thick_out(t,x,y):
    for i in range(x):
        for j in range(y):
            if i == 0:
                Set_cell_border(t.cell(0,j),top={"sz": 15, "val": "single", "color": "#000000", "space": "0"})
            if i == x-1:
                Set_cell_border(t.cell(x-1,j),bottom={"sz": 15, "val": "single", "color": "#000000", "space": "0"})
            if j == 0:
                Set_cell_border(t.cell(i,0),start={"sz": 15, "val": "single", "color": "#000000", "space": "0"})
            if j == y-1:
                Set_cell_border(t.cell(i,y-1),end={"sz": 15, "val": "single", "color": "#000000", "space": "0"})

class Main(QMainWindow, ui.Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.Benter.clicked.connect(self.EnterEvent)
        self.Bclear.clicked.connect(self.testEvent)

    def ClearEvent(self):
        self.ActivityName.setText('')
        self.Use.setText('')
        self.Department.setText('')
        self.Price.setText('')
        for x in range(48):
                self.List.itemAt(x).widget().setText('')
    def testEvent(self):
        self.ActivityName.setText('00測試')
        self.Use.setText('測試用')
        self.Department.setText('測試')
        self.Price.setText('1000')
        for x in range(48):
                self.List.itemAt(x).widget().setText('1')

    def EnterEvent(self):
        searchData = [self.ActivityName.text(), self.Department.text(), self.Price.text(), self.Use.text()]
        #print(searchData)

        #產生文件 open the docx
        doc = Document()
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
        section = doc.sections[0]
        section.bottom_margin=Cm(0.27)
        #設定基礎表格
        p1 = doc.add_paragraph()
        p1.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        title = p1.add_run('國立臺灣師範大學　 高屏地區同鄉校友會　  黏貼憑證用紙')
        title.bold = True
        title.underline = True
        title.font.size = Pt(16)

        p2 = doc.add_paragraph()
        p2.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        activityName = p2.add_run('活動名稱')
        activityName.bold = False
        activityName.underline = False
        activityName.font.size = Pt(13)
        space = '＿'*(9-len(self.ActivityName.text()))
        ActivityName = p2.add_run('＿＿'+self.ActivityName.text()+space).font.size=Pt(13)
        ActivityName.underline = True

        table = doc.add_table(rows=4, cols=11,style = 'Table Grid')
        thick_out(table,4,11)

        widths = (Cm(1.44), Cm(1.44), Cm(0.8), Cm(0.8), Cm(0.8), Cm(0.8), Cm(0.8), Cm(0.8), Cm(0.8), Cm(0.8),  Cm(7.84)) 
        for row in table.rows:
            for c,w in enumerate(widths):
                row.cells[c].width = w
        
        table.rows[0].height = Cm(0.63)
        table.rows[1].height = Cm(0.63)
        table.rows[2].height = Cm(1.875)
        table.rows[3].height = Cm(1.875)

        table.cell(0,0).merge(table.cell(0,1))
        input_word(table, 0, 0, '憑 證 編 號',1)
        input_word(table, 1, 0, '字',1)
        input_word(table, 1, 1, '號',1)
        table.cell(0,2).merge(table.cell(0,9))
        input_word(table, 0, 2, '金額',1)
        price_list = ['十', '萬', '千', '百', '十', '元', '角', '分']
        for i in range(8):
            input_word(table, 1, 2+i,price_list[i], 1)
        table.cell(0,10).merge(table.cell(1,10))
        input_word(table, 0, 10, '用途說明',1)

        #字號金額
        table.cell(2,0).merge(table.cell(3,0))
        if self.Department.text()!='':
            input_word(table, 2, 0, self.Department.text()[0],1)
        else:
            input_word(table, 2, 0, self.Department.text(),1)
        table.cell(2,1).merge(table.cell(3,1))
        if self.Price.text()!='':
            price = int(self.Price.text())
        else:
            price = 0
        for i in range(6):
            if price!=0:
                table.cell(2,7-i).merge(table.cell(3,7-i))
                input_word(table, 2, 7-i, str(price%10),1)
                price //= 10

        #用途說明
        table.cell(2,10).merge(table.cell(3,10))
        pt = table.cell(3,10).paragraphs[0]
        pt.add_run(self.Use.text()+'\n').font.size = Pt(13)

        for i in range(1,4):
            name = 'name_'+str(i)
            if self.findChild(QLineEdit, name).text()!='':
                if i!=1:
                    table.cell(3,10).paragraphs[0].add_run('\n')
                p = 'p_'+str(i)
                n = 'n_'+str(i)
                u = 'u_'+str(i)
                if int(self.findChild(QLineEdit, p).text())<0:
                    table.cell(3,10).paragraphs[0].add_run(self.findChild(QLineEdit, name).text() + '　　' + self.findChild(QLineEdit, p).text() + '元').font.size = Pt(12)
                else:
                    if self.findChild(QLineEdit, p).text()!='' and self.findChild(QLineEdit, n).text()!='':
                        total = int(self.findChild(QLineEdit, p).text())*int(self.findChild(QLineEdit, n).text())
                    else:
                        total = 0
                    table.cell(3,10).paragraphs[0].add_run(self.findChild(QLineEdit, name).text()+'\n').font.size = Pt(12)
                    table.cell(3,10).paragraphs[0].add_run(self.findChild(QLineEdit, p).text()+'(元) * '+self.findChild(QLineEdit, n).text()+'('+self.findChild(QLineEdit, u).text()+') = '+ str( total) + '元').font.size = Pt(12)
            else:
                if i!=1:
                    table.cell(3,10).add_paragraph().add_run('總計'+self.Price.text()+'元')
                    table.cell(3,10).paragraphs[1].paragraph_format.alignment=2
                break
        #簽名區
        doc.add_paragraph().add_run('').font.size=Pt(16)
        sign = doc.add_table(rows=2, cols=5,style = 'Table Grid')
        thick_out(sign,2,5)
        widths = (Cm(3), Cm(2.86), Cm(3.17), Cm(2.54), Cm(5.62)) 
        for row in sign.rows:
            for c,w in enumerate(widths):
                row.cells[c].width = w
        sign.rows[0].height = Cm(0.89)
        sign.rows[1].height = Cm(1.6)
        sign_list=['社團負責人','會計','組(股)長','經手人','查帳紀錄',]
        for i in range(5):
            input_word(sign, 0, i, sign_list[i], 4)

        p3 = doc.add_paragraph()
        p3.add_run('\n\n').font.size=Pt(16)

        table2 = doc.add_table(rows=1, cols=4,style = 'Normal Table')
        widths = (Cm(5.5), Cm(6), Cm(3.2), Cm(2)) 
        for row in table2.rows:
            for c,w in enumerate(widths):
                row.cells[c].width = w
        table2.rows[0].height = Cm(10)

        if self.name_3.text()!='':
            if self.name_4.text()=='':
                table2.cell(0,0).paragraphs[0].add_run('\n\n\n').font.size = Pt(12)
                table2.cell(0,0).add_paragraph().add_run('總計'+self.Price.text()+'元').font.size = Pt(12)
                table2.cell(0,0).paragraphs[1].paragraph_format.alignment=2
            else:
                for i in range(4,13):
                    name = 'name_'+str(i)
                    if self.findChild(QLineEdit, name).text()!='':
                        if i!=4:
                            table2.cell(0,0).paragraphs[0].add_run('\n')
                        p = 'p_'+str(i)
                        n = 'n_'+str(i)
                        u = 'u_'+str(i)
                        if int(self.findChild(QLineEdit, p).text())<0:
                            table2.cell(0,0).paragraphs[0].add_run(self.findChild(QLineEdit, name).text() + '　　' + self.findChild(QLineEdit, p).text() + '元').font.size = Pt(12)
                        else:
                            if self.findChild(QLineEdit, p).text()!='' and self.findChild(QLineEdit, n).text()!='':
                                total = int(self.findChild(QLineEdit, p).text())*int(self.findChild(QLineEdit, n).text())
                            else:
                                total = 0
                            table2.cell(0,0).paragraphs[0].add_run(self.findChild(QLineEdit, name).text()+'\n').font.size = Pt(12)
                            table2.cell(0,0).paragraphs[0].add_run(self.findChild(QLineEdit, p).text()+'(元) * '+self.findChild(QLineEdit, n).text()+'('+self.findChild(QLineEdit, u).text()+') = '+ str( total) + '元').font.size = Pt(12)
                    else:
                        if i!=4:
                            table2.cell(0,0).add_paragraph().add_run('總計'+self.Price.text()+'元').font.size = Pt(12)
                            table2.cell(0,0).paragraphs[1].paragraph_format.alignment=2
                        break
                    if i==12:
                        table2.cell(0,0).add_paragraph().add_run('總計'+self.Price.text()+'元').font.size = Pt(12)
                        table2.cell(0,0).paragraphs[1].paragraph_format.alignment=2 
        
        table3 = table2.cell(0,1).add_table(rows=13, cols=2)
        table3.style = 'Table Grid'
        widths = (Cm(5.3), Cm(0.7)) 
        for row in table3.rows:
            for c,w in enumerate(widths):
                row.cells[c].width = w
        table3.rows[0].height = Cm(1.5)
        for i in range(1,13):
            table3.rows[i].height = Cm(0.7)
        input_word(table3, 0, 0, '附　　　件',1)
        table3.cell(0,0).merge(table3.cell(0,1))
        thick_out(table3,13,2)        
        table3_list = ['請購單','請修單','監驗','驗收','估價單','圖說','樣張','電文','印模','驗收報告','','其他文件']
        for i in range(12):
            input_word(table3, 1+i, 0, table3_list[i],0)
            Set_cell_border(table3.cell(1+i,0),end={"sz": 12, "val": "none", "color": "#000000", "space": "0"})
            if not 2<=i<=3:
                input_word(table3, 1+i, 1, '張',2)
            Set_cell_border(table3.cell(1+i,1),start={"sz": 12, "val": "none", "color": "#000000", "space": "0"})
            if i==7:
                Set_cell_border(table3.cell(1+i,0),bottom={"sz": 4, "val": "double", "color": "#000000", "space": "0"})
                Set_cell_border(table3.cell(1+i,1),bottom={"sz": 4, "val": "double", "color": "#000000", "space": "0"})

        
        table4 = table2.cell(0,3).add_table(rows=6, cols=1)
        table4.style = 'Table Grid'
        for row in table4.rows:
            row.cells[0].width = Cm(2)
        for i in range(3):
            table4.rows[0+2*i].height = Cm(0.85)
            table4.rows[1+2*i].height = Cm(2.5)
        thick_out(table4,6,1)
        table4_list = ['監 驗','驗 收','保 管']
        for i in range(3):
            input_word(table4, 0+2*i, 0, table4_list[i], 1)

        #存檔
        ad = os.path.abspath(os.path.dirname(__file__)).replace( '\\','/')
        doc.save(ad+'/黏憑.docx')
        
        convert(ad+'/黏憑.docx',ad+'/黏憑.pdf')

if __name__ == '__main__':
    import sys
    app = QtWidgets.QApplication(sys.argv)
    window = Main()
    window.show()
    sys.exit(app.exec_())